#!/usr/bin/env python3
"""Convertit Strategie-Canaria-Particuliers.md en .docx mis en page."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = RGBColor(0xEC, 0x69, 0x25)
PETROL = RGBColor(0x1F, 0x3A, 0x42)
GREY = RGBColor(0x55, 0x55, 0x55)

SRC = "Strategie-Canaria-Particuliers.md"
OUT = "Strategie-Canaria-Particuliers.docx"

doc = Document()

# --- Base styles ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, (size, color) in {
    "Heading 1": (16, ORANGE),
    "Heading 2": (13, PETROL),
    "Heading 3": (11.5, PETROL),
}.items():
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_runs(paragraph, text):
    """Gère **gras** et `code` inline."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = PETROL
        else:
            paragraph.add_run(part)


def add_table(lines):
    # lines: list of markdown table rows (incl. header + separator)
    rows = [
        [c.strip() for c in ln.strip().strip("|").split("|")]
        for ln in lines
        if ln.strip()
    ]
    header = rows[0]
    body = [r for r in rows[2:]]  # skip separator row
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 2"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, txt in enumerate(header):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, txt)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "EC6925")
    for r in body:
        cells = table.add_row().cells
        for i in range(ncols):
            txt = r[i] if i < len(r) else ""
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], txt)
    doc.add_paragraph()


with open(SRC, encoding="utf-8") as f:
    raw = f.read().split("\n")

i = 0
while i < len(raw):
    line = raw[i]
    stripped = line.strip()

    # Table block
    if stripped.startswith("|") and i + 1 < len(raw) and set(raw[i + 1].strip()) <= set("|-: "):
        block = []
        while i < len(raw) and raw[i].strip().startswith("|"):
            block.append(raw[i])
            i += 1
        add_table(block)
        continue

    if not stripped:
        i += 1
        continue

    # Horizontal rule
    if stripped == "---":
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "EC6925")
        pbdr.append(bottom)
        pPr.append(pbdr)
        i += 1
        continue

    # Headings
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        h = doc.add_heading(stripped[2:], level=1)
    # Blockquote
    elif stripped.startswith(">"):
        txt = stripped.lstrip("> ").strip()
        txt = txt.replace("⚠️", "").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "EC6925")
        pbdr.append(left)
        pPr.append(pbdr)
        add_runs(p, txt)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = GREY
    # Checklist
    elif stripped.startswith("- [ ]"):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, "☐ " + stripped[5:].strip())
    # Bullets
    elif stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, stripped[2:])
    # Numbered
    elif re.match(r"^\d+\.\s", stripped):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
    # Italic-only line (subtitle / date)
    elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
        p = doc.add_paragraph()
        r = p.add_run(stripped.strip("*"))
        r.italic = True
        r.font.color.rgb = GREY
    else:
        p = doc.add_paragraph()
        add_runs(p, stripped)
    i += 1

doc.save(OUT)
print("OK ->", OUT)
